"""Build a DOCX version of the Alignerr CV."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "cv-alignerr-ai-code-ranking.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

NAVY = RGBColor(0x0B, 0x2B, 0x4E)
GREY = RGBColor(0x44, 0x4B, 0x54)


def add_heading(text, size=13, color=NAVY, space_before=10, space_after=2, rule=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if rule:
        rule_p = doc.add_paragraph()
        rule_p.paragraph_format.space_before = Pt(0)
        rule_p.paragraph_format.space_after = Pt(2)
        rule_run = rule_p.add_run("_" * 95)
        rule_run.font.color.rgb = NAVY
        rule_run.font.size = Pt(6)


def add_role(title, org, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{title}")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(f"  —  {org}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r3 = p2.add_run(dates)
    r3.italic = True
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = GREY


def add_body(text, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def add_bullet(segments, space_after=2):
    """segments: list of (text, bold)."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10.5)


# ---------- Header ----------
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(0)
name_run = name_p.add_run("ONIS EMEM")
name_run.bold = True
name_run.font.size = Pt(22)
name_run.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(0)
sub_run = sub_p.add_run("Senior Software Engineer  |  AI Code Evaluation & Productization")
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = GREY

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(2)
contact_run = contact_p.add_run(
    "Tallinn, Estonia  ·  onis.c.emem@gmail.com  ·  +372 589 87 957"
)
contact_run.font.size = Pt(10)

links_p = doc.add_paragraph()
links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
links_p.paragraph_format.space_after = Pt(4)
links_run = links_p.add_run(
    "linkedin.com/in/onis-emem  ·  onis-emem.com  ·  github.com/oniso20"
)
links_run.font.size = Pt(10)
links_run.font.color.rgb = GREY

# ---------- Summary ----------
add_heading("Summary")
add_body(
    "Senior engineer with 7+ years shipping production code across full-stack web, payments, "
    "and AI systems, now operating as a PM who still writes and reviews code daily. Deep "
    "experience with LLM evaluation in production — built eval-gated AI features using Langfuse "
    "and LLM-as-judge, defined response-quality rubrics, and closed feedback loops that cut "
    "hallucinations. Fluent in JavaScript/TypeScript, Python, and Node.js; comfortable reading "
    "and reviewing Java, Go, and C-family code. Known for small PRs, sharp code review, and "
    "articulating exactly what makes one solution better than another."
)

# ---------- Relevant Strengths ----------
add_heading("Relevant Strengths for AI Code Ranking")
add_bullet([
    ("Code review at scale: ", True),
    ("Review PRs across Next.js, Node, and Python stacks for Inktrail, TechSynergy, and "
     "Laferla. Consistently catch logic errors, missing edge cases, security issues (auth, "
     "injection, secret handling), and readability regressions before merge.", False),
])
add_bullet([
    ("LLM output evaluation: ", True),
    ("Shipped Langfuse + LLM-as-judge on Hub AI; wrote rubrics for correctness, factuality, "
     "and tone; +20% response quality in production. Round-trip server-side validation "
     "errors back to the model so it self-corrects in 1–2 turns.", False),
])
add_bullet([
    ("Multi-language fluency: ", True),
    ("JavaScript/TypeScript, Python, Node.js daily. Java/C++ from academic work and "
     "competitive problem solving; Go and Rust read-fluent for review.", False),
])
add_bullet([
    ("Algorithms & complexity: ", True),
    ("M.Sc. Industrial Ecology with ML modelling (supervised learning on PV waste "
     "forecasting, ACS-published). Built real-time tracking, funnel analytics, and "
     "cost-reconciliation pipelines from scratch.", False),
])
add_bullet([
    ("RLHF & preference data awareness: ", True),
    ("Productized LLM features that depend on preference signals; designed feedback capture "
     "and evaluation loops that inform model- and prompt-level improvements.", False),
])

# ---------- Core Skills ----------
add_heading("Core Skills")
add_bullet([
    ("Languages: ", True),
    ("JavaScript, TypeScript, Python, Node.js, SQL  ·  Reading: Java, Go, C++, Rust", False),
])
add_bullet([
    ("AI / ML: ", True),
    ("LLM evaluation (RAG, LLM-as-judge), prompt engineering, Langfuse, eval-gated release, "
     "cost reconciliation, RLHF concepts", False),
])
add_bullet([
    ("Engineering: ", True),
    ("React, Next.js, Express, REST/GraphQL API design, BFF patterns, Azure/AWS, Git, "
     "CI/CD, testing & TDD, performance profiling", False),
])
add_bullet([
    ("Quality: ", True),
    ("Code review, design patterns, refactoring, security (OWASP, secret hygiene, input "
     "validation), observability", False),
])
add_bullet([
    ("Domains: ", True),
    ("Payments (Wise), insurance, iGaming integrations, KYC/KYB, marketplaces", False),
])

# ---------- Experience ----------
add_heading("Experience")

add_role(
    "Senior Product Manager (with hands-on AI/eng ownership)",
    "Hub88 / Yolo Group · Tallinn",
    "2023 – Present",
)
add_body(
    "Own HubConnect (B2B integration platform) and Hub AI (internal LLM assistant). Still in "
    "the codebase — review PRs, write specs at code level, pair on integration and evaluation "
    "logic."
)
add_bullet([
    ("Hub AI — eval-gated LLM productization. ", True),
    ("Built evaluation harness with Langfuse + LLM-as-judge before widening access. Defined "
     "response-quality rubric (correctness, groundedness, tone), shipped fast feedback loops, "
     "improved response quality by 20% and reduced hallucinations in production. 100+ users, "
     "350+ sessions in six weeks.", False),
])
add_bullet([
    ("HubConnect — engineering reviewer for integration code. ", True),
    ("Review API contracts, schema migrations, and integration code between suppliers and "
     "operators. Catch edge cases, auth and replay issues, and idempotency bugs before rollout.",
     False),
])
add_bullet([
    ("1-Click Onboarding — workflow replatforming. ", True),
    ("Replaced manual onboarding with a self-service flow; +80% partners integrating within "
     "two weeks, +140% production-ready integrations the following quarter.", False),
])
add_bullet([
    ("Promotions marketplace. ", True),
    ("Shipped two-sided marketplace — 60% supplier adoption, 70% operator opt-in in month one.",
     False),
])

add_role(
    "Product Manager & Engineer",
    "Laferla Insurance Group · Tallinn",
    "2022 – Present",
)
add_body(
    "Rebuilt a 40-year-old insurer's customer platform on Next.js + Node with a BFF to the "
    "internal CRM. Hands-on engineering lead on quoting, KYC/KYB, and admin tooling."
)
add_bullet([("Rewrote quote engine: 3–7 days → instant turnaround while preserving underwriting rules.", False)])
add_bullet([("KYC automation: days → 1–2 minutes (SumSub integration, re-engineered review).", False)])
add_bullet([("KYB: 3–7 days → 24 hours.", False)])
add_bullet([("Admin BFF + Azure SSO: −40% internal workload, +50% NPS.", False)])
add_bullet([("Owned code review, release process, and on-call for the platform.", False)])

add_role(
    "Software Developer & Scrum Master",
    "100Devs · Boston (Remote)",
    "2021 – 2022",
)
add_body(
    "Shipped full-stack React/Node.js apps for multiple clients. Translated ambiguous goals "
    "into MVPs; −30% time-to-market via tight feedback loops and code-review-led development."
)

add_role(
    "Payment Operations Specialist",
    "Wise (TransferWise) · Tallinn",
    "2021 – 2022",
)
add_body(
    "Analysed large transaction datasets (SQL/Python) to find refund-logic bottlenecks; "
    "partnered with Treasury, Compliance, and Product engineers to implement fixes. +32% "
    "workflow throughput."
)

add_role(
    "Software Developer & Scrum Master",
    "Dado Food · Enugu, Nigeria (Remote)",
    "2019 – 2021",
)
add_body(
    "Built real-time order tracking on Node.js — −15% support inquiries. Frontend and asset "
    "optimisation: +30% app performance."
)

add_role(
    "Python Data Scientist / ML Researcher",
    "ITMO University × University of Southern Denmark",
    "2018 – 2025 (fractional post-grad)",
)
add_body(
    "Co-authored ACS-published research using supervised ML to forecast solar PV waste across "
    "15 West African countries. Python, scikit-learn, time-series modelling, uncertainty "
    "quantification. Publication: Environmental Science & Technology, 2025."
)

# ---------- Side Projects ----------
add_heading("Side Projects & Open Source")
add_bullet([
    ("Inktrail (inktrail.ai) — ", True),
    ("Collaborative workspace (docs, canvas, publishing) experimenting with AI + human "
     "co-editing. Real-time sync, per-call cost reconciliation, eval-gated AI features, "
     "server-side validation round-tripped back to the model for self-correction. Solo-built.",
     False),
])
add_bullet([
    ("oniso20/conductor — ", True),
    ("Open-source CLI for running multi-stage Claude agent workflows.", False),
])
add_bullet([
    ("TechSynergy (Co-founder, CPO) — ", True),
    ("Africa-first hiring & payments platform. AI-assisted talent matching, multi-currency "
     "wallets. €7K MRR, €100K+ revenue, 50+ projects shipped, 900+ talent community.", False),
])

# ---------- Education ----------
add_heading("Education")
add_bullet([("Vocational Diploma, ICT — Helsinki Business College, 2022–2023", False)])
add_bullet([("M.Sc. Exchange, Life Cycle Engineering — University of Southern Denmark, 2019–2020", False)])
add_bullet([("M.Sc. Industrial Ecology — ITMO University, 2018–2020", False)])
add_bullet([("B.Tech. Environmental Engineering — Rivers State University, 2004–2010", False)])

# ---------- Certifications ----------
add_heading("Certifications")
add_bullet([("Software Architecture — Global Dev Experts, 2024", False)])
add_bullet([("Agile Project & Delivery Management — ICAgile, 2023", False)])
add_bullet([("Power BI Data Analyst Associate — Microsoft, 2021", False)])

doc.save(OUT)
print(f"Wrote {OUT}")
